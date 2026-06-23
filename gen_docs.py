import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 确保 txt 目录存在
out_dir = Path("txt")
out_dir.mkdir(exist_ok=True)

# 10 个源文件
files = [
    "day01_first_chat.py",
    "day02_control_output.py",
    "day03_structured_output.py",
    "day04_memory_chat.py",
    "day05_tool_calling.py",
    "day06_chatbot_project.py",
    "day07_rag_load_split.py",
    "day08_rag_embed_retrieve.py",
    "day09_minimal_rag.py",
    "day10_llm_principles.py",
]

for fname in files:
    src = Path(fname)
    if not src.exists():
        print(f"跳过（不存在）: {fname}")
        continue

    content = src.read_text(encoding="utf-8")

    doc = Document()

    # 标题：文件名
    title = doc.add_heading(src.stem, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 代码内容用等宽字体
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # 保存
    out_path = out_dir / src.with_suffix(".docx").name
    doc.save(str(out_path))
    print(f"已生成: {out_path}")

print("\n全部完成！")
